from __future__ import annotations

import re
import sys
import urllib.request
from html import unescape
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


SOURCE_URL = "https://ffmpeg.org/ffmpeg-all.html"
OUT_PATH = Path("Tai_lieu_FFmpeg_chi_tiet_tieng_Viet.docx")


class TocParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.in_contents = False
        self.in_anchor = False
        self.depth = 0
        self.current: list[str] = []
        self.items: list[tuple[int, str]] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = dict(attrs)
        if tag == "div" and attrs_dict.get("class") == "contents":
            self.in_contents = True
            self.depth = 1
            return
        if not self.in_contents:
            return
        if tag == "div":
            self.depth += 1
        if tag == "a" and attrs_dict.get("id", "").startswith("toc-"):
            self.in_anchor = True
            self.current = []

    def handle_endtag(self, tag: str) -> None:
        if not self.in_contents:
            return
        if tag == "a" and self.in_anchor:
            text = clean_text("".join(self.current))
            level = toc_level(text)
            if text and level > 0:
                self.items.append((level, text))
            self.in_anchor = False
            self.current = []
        if tag == "div":
            self.depth -= 1
            if self.depth <= 0:
                self.in_contents = False

    def handle_data(self, data: str) -> None:
        if self.in_contents and self.in_anchor:
            self.current.append(data)


def clean_text(value: str) -> str:
    value = unescape(value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def toc_level(text: str) -> int:
    match = re.match(r"^(\d+(?:\.\d+)*)\s+", text)
    if not match:
        return 0
    return match.group(1).count(".") + 1


def title_without_number(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*\s+", "", text).strip()


def fetch_toc() -> list[tuple[int, str]]:
    with urllib.request.urlopen(SOURCE_URL, timeout=30) as response:
        html = response.read().decode("utf-8", errors="replace")
    parser = TocParser()
    parser.feed(html)
    return parser.items


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def section_notes(title: str) -> tuple[str, list[str], str]:
    key = title.lower()
    if "synopsis" in key:
        return (
            "Cú pháp tổng quát của ffmpeg: tùy chọn chung, một hoặc nhiều input, rồi một hoặc nhiều output. Vị trí option rất quan trọng.",
            [
                "Option đặt trước `-i` thường áp dụng cho input kế tiếp.",
                "Option đặt trước output thường áp dụng cho output đó.",
                "Một lệnh ffmpeg có thể đọc nhiều input và ghi nhiều output trong cùng lần chạy.",
            ],
            "ffmpeg -i input.mp4 output.webm",
        )
    if "description" == key:
        return (
            "FFmpeg là bộ công cụ xử lý multimedia: đọc, giải mã, lọc, mã hóa, đóng gói và ghi ra file/stream.",
            [
                "Input có thể là file, URL, pipe, camera, microphone hoặc stream mạng.",
                "Output có thể là file MP4, HLS, MPEG-TS, RTMP, SRT, image sequence và nhiều dạng khác.",
                "Khi không chỉ định rõ, ffmpeg cố tự chọn stream, codec và muxer phù hợp.",
            ],
            "ffmpeg -i in.mov -c:v libx264 -c:a aac out.mp4",
        )
    if "streamcopy" in key:
        return (
            "Streamcopy là sao chép luồng audio/video/subtitle mà không giải mã và không mã hóa lại.",
            [
                "Rất nhanh vì chỉ remux hoặc cắt ở cấp container.",
                "Dùng `-c copy` khi codec gốc đã phù hợp với output.",
                "Không dùng được nếu bạn cần resize, đổi bitrate, burn subtitle hoặc áp filter.",
            ],
            "ffmpeg -i input.mkv -c copy output.mp4",
        )
    if "transcoding" in key:
        return (
            "Transcoding là quá trình decode stream gốc rồi encode lại sang codec/tham số khác.",
            [
                "Cần khi đổi codec, giảm dung lượng, đổi resolution, đổi sample rate hoặc tối ưu tương thích.",
                "Tốn CPU/GPU hơn streamcopy và có thể làm giảm chất lượng nếu encode lossy.",
                "Các option thường gặp: `-c:v`, `-c:a`, `-b:v`, `-crf`, `-preset`, `-pix_fmt`.",
            ],
            "ffmpeg -i input.mov -c:v libx264 -crf 23 -preset medium -c:a aac output.mp4",
        )
    if "filter" in key:
        return (
            "Filter là các bước biến đổi audio/video sau khi decode và trước khi encode.",
            [
                "Simple filtergraph dùng `-vf` hoặc `-af` cho một luồng đơn giản.",
                "Complex filtergraph dùng `-filter_complex` khi có nhiều input/output hoặc nối nhiều nhánh.",
                "Filter phổ biến: `scale`, `crop`, `fps`, `overlay`, `drawtext`, `aresample`, `loudnorm`.",
            ],
            "ffmpeg -i input.mp4 -vf scale=1280:-2 -c:v libx264 output.mp4",
        )
    if "stream selection" in key or "stream specifier" in key or "map" in key:
        return (
            "Nhóm này quyết định stream nào từ input sẽ đi vào output.",
            [
                "`-map` giúp chọn thủ công video/audio/subtitle thay vì để ffmpeg tự chọn.",
                "Specifier như `0:v:0`, `0:a:1` nghĩa là input 0, video/audio, stream thứ mấy.",
                "Rất quan trọng khi file có nhiều audio track, subtitle hoặc nhiều camera angle.",
            ],
            "ffmpeg -i input.mkv -map 0:v:0 -map 0:a:1 -c copy output.mp4",
        )
    if "option" in key or "avoptions" in key or "preset" in key:
        return (
            "Options là hệ thống tham số điều khiển hành vi của ffmpeg, codec, muxer, demuxer, filter và protocol.",
            [
                "Generic options dùng chung như `-h`, `-version`, `-formats`, `-codecs`.",
                "Main options xử lý input/output như `-i`, `-y`, `-ss`, `-t`, `-to`.",
                "Codec AVOptions phụ thuộc codec cụ thể, nên nên tra bằng `ffmpeg -h encoder=...`.",
            ],
            "ffmpeg -h encoder=libx264",
        )
    if "example" in key:
        return (
            "Phần ví dụ giúp nối khái niệm với lệnh thực tế: convert file, record thiết bị, grab màn hình hoặc tạo output thử nghiệm.",
            [
                "Hãy đọc ví dụ theo thứ tự input -> xử lý -> output.",
                "Khi copy ví dụ, cần đổi tên thiết bị, path và codec theo máy của bạn.",
                "Với Windows, tên thiết bị capture thường xem bằng `ffmpeg -list_devices true -f dshow -i dummy`.",
            ],
            "ffmpeg -f lavfi -i testsrc2=size=1280x720:rate=30 -t 5 test.mp4",
        )
    if "syntax" in key or "quoting" in key or "date" in key or "duration" in key or "ratio" in key or "color" in key or "channel" in key:
        return (
            "Syntax mô tả cách viết giá trị trong lệnh: thời lượng, kích thước, tỉ lệ, màu, channel layout và escaping.",
            [
                "Thời lượng có thể viết `10`, `00:00:10`, `1.5`, hoặc dạng có đơn vị tùy ngữ cảnh.",
                "Video size thường là `1920x1080`, `1280x720`, hoặc alias như `hd720`.",
                "Filter phức tạp trên Windows cần chú ý dấu nháy kép và escape ký tự đặc biệt.",
            ],
            "ffmpeg -ss 00:01:00 -t 10 -i input.mp4 -c copy clip.mp4",
        )
    if "expression" in key:
        return (
            "Expression Evaluation là ngôn ngữ biểu thức dùng trong filter và một số option.",
            [
                "Có biến như `w`, `h`, `iw`, `ih`, `t`, `n` tùy filter.",
                "Dùng để tính động kích thước, vị trí overlay, crop, volume theo thời gian.",
                "Biểu thức sai thường làm filter báo lỗi khó đọc, nên test từng phần nhỏ.",
            ],
            "ffmpeg -i input.mp4 -vf \"scale=iw/2:ih/2\" small.mp4",
        )
    if "decoder" in key:
        return (
            "Decoder đọc dữ liệu đã nén và biến thành frame/audio sample thô để FFmpeg xử lý tiếp.",
            [
                "Thường ffmpeg tự chọn decoder theo codec của input.",
                "Có thể chỉ định decoder khi cần phần cứng hoặc thư viện cụ thể.",
                "Decoder không quyết định định dạng file output; encoder và muxer mới quyết định phần đó.",
            ],
            "ffmpeg -decoders",
        )
    if "encoder" in key:
        return (
            "Encoder biến frame/audio sample thô thành dữ liệu nén như H.264, HEVC, AAC, Opus hoặc AV1.",
            [
                "`libx264` phổ biến cho H.264, `libx265` cho HEVC, `aac` cho audio AAC.",
                "Chất lượng và dung lượng thường điều khiển bằng CRF/bitrate/preset.",
                "Encoder có thể phụ thuộc cách FFmpeg được build, không phải bản nào cũng có mọi encoder.",
            ],
            "ffmpeg -i input.mp4 -c:v libx264 -crf 22 -c:a aac output.mp4",
        )
    if "hardware" in key or "qsv" in key or "vaapi" in key or "cuda" in key or "vulkan" in key or "videotoolbox" in key:
        return (
            "Nhóm phần cứng dùng GPU/ASIC để tăng tốc decode, filter hoặc encode.",
            [
                "Ưu điểm là nhanh và tiết kiệm CPU.",
                "Nhược điểm là setup phức tạp hơn, chất lượng/option khác encoder phần mềm.",
                "Cần kiểm tra driver, build FFmpeg và thiết bị hỗ trợ.",
            ],
            "ffmpeg -hwaccels",
        )
    if "muxer" in key:
        return (
            "Muxer đóng gói các stream đã encode/copy vào container hoặc format output.",
            [
                "Ví dụ muxer: MP4, Matroska, MPEG-TS, HLS, DASH, image2.",
                "Muxer quyết định đuôi file, metadata container, cách chia segment và quy tắc stream hợp lệ.",
                "Một codec có thể nằm trong nhiều container, nhưng không phải container nào cũng chứa mọi codec.",
            ],
            "ffmpeg -i input.mp4 -c copy -f hls playlist.m3u8",
        )
    if "demuxer" in key:
        return (
            "Demuxer đọc container/input format và tách ra các stream riêng như video, audio, subtitle.",
            [
                "Ví dụ demuxer: mov/mp4, matroska, hls, mpegts, image2, concat.",
                "Demuxer làm việc trước decoder.",
                "Có thể ép demuxer bằng `-f` khi input không có đuôi hoặc là pipe.",
            ],
            "ffmpeg -f concat -safe 0 -i list.txt -c copy joined.mp4",
        )
    if "protocol" in key or key in {"http", "https", "rtmp", "rtsp", "srt", "tcp", "udp", "file", "pipe"}:
        return (
            "Protocol là lớp truy cập dữ liệu: file local, HTTP(S), RTMP, RTSP, SRT, TCP/UDP, pipe và các dạng cache.",
            [
                "Protocol ảnh hưởng timeout, reconnect, latency, buffer và cách đọc/ghi stream.",
                "Với live stream, protocol quan trọng không kém codec.",
                "Nhiều protocol có option riêng viết trong URL hoặc bằng AVOption.",
            ],
            "ffmpeg -re -i input.mp4 -f flv rtmp://server/app/key",
        )
    if "filter" in key:
        return (
            "Filters biến đổi hình ảnh hoặc âm thanh ở giữa pipeline.",
            [
                "Video filter xử lý frame: scale, crop, overlay, drawtext, fps, format.",
                "Audio filter xử lý sample: volume, atrim, loudnorm, aresample, amix.",
                "Filtergraph càng phức tạp càng nên đặt trong file script hoặc viết từng bước để debug.",
            ],
            "ffmpeg -i input.mp4 -filter_complex \"[0:v]scale=1280:-2[v]\" -map \"[v]\" output.mp4",
        )
    if "device" in key:
        return (
            "Device là input/output đặc biệt như camera, microphone, màn hình, audio device hoặc capture card.",
            [
                "Tên device phụ thuộc hệ điều hành: dshow trên Windows, avfoundation trên macOS, v4l2/x11grab trên Linux.",
                "Luôn liệt kê device trước khi record.",
                "Record device thường cần chỉ định format, resolution, framerate hoặc audio sample rate.",
            ],
            "ffmpeg -list_devices true -f dshow -i dummy",
        )
    if "scaler" in key or "resampler" in key or "swscale" in key or "swresample" in key:
        return (
            "Scaler/resampler là lớp chuyển đổi kích thước, pixel format, sample rate và channel layout.",
            [
                "Video scaling ảnh hưởng độ nét, tốc độ và khả năng tương thích.",
                "Audio resampling đổi tần số mẫu, số kênh hoặc định dạng sample.",
                "Thường được gọi gián tiếp qua filter hoặc encoder yêu cầu format cụ thể.",
            ],
            "ffmpeg -i input.wav -ar 48000 -ac 2 output.wav",
        )
    if "bitstream" in key:
        return (
            "Bitstream filter sửa hoặc chuyển cấu trúc bitstream mà không decode toàn bộ frame.",
            [
                "Hay dùng khi remux H.264/H.265 giữa MP4, Annex B, MPEG-TS.",
                "Nhanh hơn transcode vì không encode lại.",
                "Cần khi container/protocol yêu cầu layout bitstream khác.",
            ],
            "ffmpeg -i input.mp4 -c copy -bsf:v h264_mp4toannexb output.ts",
        )
    if "metadata" in key:
        return (
            "Metadata là thông tin mô tả file/stream như title, language, creation_time, artist hoặc tags riêng.",
            [
                "Có metadata cấp container và cấp từng stream.",
                "`-map_metadata` điều khiển copy metadata giữa input/output.",
                "`-metadata` ghi hoặc thay tag cụ thể.",
            ],
            "ffmpeg -i input.mp4 -metadata title=\"Demo\" -c copy output.mp4",
        )
    return (
        f"Mục `{title}` là một phần trong tài liệu ffmpeg-all, thường mô tả một thành phần, nhóm option hoặc hành vi cụ thể trong pipeline FFmpeg.",
        [
            "Khi học mục này, hãy xác định nó thuộc giai đoạn nào: input, demux, decode, filter, encode, mux hay output.",
            "Tra option chi tiết bằng `ffmpeg -h`, `ffmpeg -h full`, hoặc `ffmpeg -h component=name` nếu FFmpeg hỗ trợ.",
            "Khi áp dụng thực tế, nên chạy thử trên file ngắn trước để tránh mất thời gian encode dài.",
        ],
        "ffmpeg -hide_banner -h",
    )


def build_doc(toc: list[tuple[int, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Tài liệu FFmpeg chi tiết bằng tiếng Việt", 0)
    doc.add_paragraph(
        "Tài liệu này được biên soạn dựa trên mục lục của trang chính thức "
        f"{SOURCE_URL}. Nội dung là phần giải thích, diễn giải và ví dụ thực hành bằng tiếng Việt."
    )
    doc.add_paragraph(
        "Mục tiêu: giúp bạn hiểu FFmpeg theo pipeline thực tế: input -> demux -> decode -> filter -> encode -> mux -> output."
    )

    doc.add_heading("Cách đọc tài liệu này", 1)
    add_bullets(
        doc,
        [
            "Nếu mới học, đọc từ mục 1 đến 9 trước: cú pháp, mô tả, streamcopy, transcoding, filter, stream selection và options.",
            "Nếu làm HLS/livestream, đọc thêm muxer/demuxer/protocol, đặc biệt HLS, MPEG-TS, RTMP, RTSP, SRT, HTTP.",
            "Nếu tối ưu chất lượng video, tập trung encoder, CRF/bitrate, pixel format, scaler và hardware acceleration.",
            "Nếu debug file, dùng ffprobe song song với ffmpeg để xem stream, codec, metadata và packet/frame.",
        ],
    )

    doc.add_heading("Bảng mục lục rút ra từ ffmpeg-all", 1)
    for level, item in toc:
        indent = "  " * max(level - 1, 0)
        doc.add_paragraph(f"{indent}{item}")
    doc.add_paragraph(f"Tổng số mục trong mục lục đã parse từ tài liệu gốc: {len(toc)}.")

    doc.add_page_break()
    doc.add_heading("Phần giải thích chi tiết theo mục lục", 1)

    explained_count = 0
    for level, raw_title in toc:
        if explained_count >= 90:
            break
        title = title_without_number(raw_title)
        if level > 3:
            continue
        heading_level = min(level + 1, 4)
        doc.add_heading(raw_title, heading_level)
        overview, bullets, example = section_notes(title)
        doc.add_paragraph(overview)
        add_bullets(doc, bullets)
        doc.add_paragraph("Ví dụ lệnh:")
        add_code(doc, example)
        explained_count += 1

    doc.add_page_break()
    doc.add_heading("Cheat sheet lệnh FFmpeg thường dùng", 1)
    cheats = [
        ("Xem thông tin file", "ffprobe -hide_banner -i input.mp4"),
        ("Đổi container không encode lại", "ffmpeg -i input.mkv -c copy output.mp4"),
        ("Encode MP4 H.264 phổ thông", "ffmpeg -i input.mov -c:v libx264 -crf 23 -preset medium -c:a aac output.mp4"),
        ("Resize giữ tỉ lệ", "ffmpeg -i input.mp4 -vf scale=1280:-2 output.mp4"),
        ("Cắt nhanh không encode", "ffmpeg -ss 00:01:00 -i input.mp4 -t 10 -c copy clip.mp4"),
        ("Cắt chính xác có encode", "ffmpeg -ss 00:01:00 -i input.mp4 -t 10 -c:v libx264 -c:a aac clip.mp4"),
        ("Tạo HLS", "ffmpeg -i input.mp4 -c:v libx264 -c:a aac -f hls -hls_time 6 -hls_list_size 0 playlist.m3u8"),
        ("Tải HLS thành MP4 nếu stream cho phép", "ffmpeg -i https://example.com/master.m3u8 -c copy output.mp4"),
        ("Chỉ lấy audio", "ffmpeg -i input.mp4 -vn -c:a mp3 output.mp3"),
        ("Burn subtitle", "ffmpeg -i input.mp4 -vf subtitles=sub.srt -c:a copy output.mp4"),
    ]
    for label, command in cheats:
        doc.add_paragraph(label, style="List Bullet")
        add_code(doc, command)

    doc.add_heading("Nguồn tham khảo", 1)
    doc.add_paragraph(f"FFmpeg official ffmpeg-all documentation: {SOURCE_URL}")
    doc.add_paragraph("FFmpeg documentation index: https://ffmpeg.org/documentation.html")
    doc.add_paragraph("FFmpeg project homepage: https://ffmpeg.org/")

    doc.save(OUT_PATH)


def main() -> int:
    toc = fetch_toc()
    if len(toc) < 50:
        print(f"TOC parse failed or too short: {len(toc)} items", file=sys.stderr)
        return 1
    build_doc(toc)
    print(f"Created {OUT_PATH.resolve()} with {len(toc)} TOC items parsed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
