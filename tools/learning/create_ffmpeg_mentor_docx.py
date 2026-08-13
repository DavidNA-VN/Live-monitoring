from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT_PATH = Path("Tai_lieu_FFmpeg_chi_tiet_tieng_Viet.docx")
SOURCE_URL = "https://ffmpeg.org/ffmpeg-all.html"


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 3"].font.name = "Arial"


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, 1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, 2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, 3)


def add_command_box(doc: Document, title: str, command: str, explain: list[str]) -> None:
    h3(doc, title)
    code(doc, command)
    bullets(doc, explain)


def build() -> None:
    doc = Document()
    setup(doc)

    title = doc.add_heading("FFmpeg từ góc nhìn mentor", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Tài liệu tiếng Việt giúp hiểu công cụ, không dịch máy documentation.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p(
        doc,
        "Tài liệu này được viết lại dựa trên tài liệu chính thức ffmpeg-all của FFmpeg, "
        "nhưng mục tiêu không phải là dịch từng dòng. Mục tiêu là giúp bạn có mô hình tư duy đúng để đọc lệnh, viết lệnh, debug lỗi và áp dụng FFmpeg vào media monitoring, HLS, video processing hoặc automation.",
    )
    p(
        doc,
        "Nếu coi FFmpeg như một chiếc hộp đen khổng lồ thì nó rất khó học. Cách dễ hơn là nhìn nó như một pipeline: đọc input, tách stream, giải mã, xử lý, mã hóa lại, đóng gói và ghi output. Khi bạn hiểu pipeline này, phần lớn option trong documentation sẽ có chỗ đứng rõ ràng.",
    )

    h1(doc, "1. FFmpeg là gì")
    p(
        doc,
        "FFmpeg là bộ công cụ dòng lệnh và thư viện để xử lý audio/video. Khi bạn chạy lệnh `ffmpeg`, bạn đang dùng một chương trình có thể đọc rất nhiều định dạng media, xử lý chúng, rồi xuất ra định dạng khác. Nó không chỉ là công cụ convert video. Nó còn dùng để cắt ghép, resize, tạo thumbnail, normalize audio, tạo HLS, nhận stream RTSP, phát RTMP/SRT, trích metadata và rất nhiều tác vụ media khác.",
    )
    bullets(
        doc,
        [
            "ffmpeg: chương trình chính để xử lý input thành output.",
            "ffprobe: công cụ xem thông tin file/stream, cực kỳ quan trọng khi debug.",
            "ffplay: trình phát media đơn giản, hữu ích để test nhanh input/output.",
            "libavcodec, libavformat, libavfilter: các thư viện nền bên dưới ffmpeg.",
        ],
    )
    p(
        doc,
        "Trong thực tế, bạn nên học `ffmpeg` và `ffprobe` song song. ffmpeg dùng để làm, ffprobe dùng để hiểu file đang có gì.",
    )

    h1(doc, "2. Mô hình pipeline quan trọng nhất")
    p(doc, "Một lệnh FFmpeg thường đi qua chuỗi sau:")
    code(doc, "Input -> Demuxer -> Decoder -> Filter -> Encoder -> Muxer -> Output")
    bullets(
        doc,
        [
            "Input là nguồn dữ liệu: file mp4, link m3u8, RTSP camera, microphone, pipe...",
            "Demuxer đọc container/format và tách ra các stream: video, audio, subtitle.",
            "Decoder giải nén stream thành dữ liệu thô: frame video hoặc sample audio.",
            "Filter xử lý dữ liệu thô: scale, crop, overlay, volume, loudnorm, drawtext...",
            "Encoder nén lại dữ liệu thành codec mới: H.264, H.265, AAC, Opus...",
            "Muxer đóng gói các stream vào output container: mp4, mkv, hls, mpegts...",
        ],
    )
    p(
        doc,
        "Sai lầm phổ biến khi mới học là lẫn lộn codec với container. MP4 không phải codec, nó là container. H.264 mới là codec video. AAC là codec audio. Một file `.mp4` có thể chứa video H.264 và audio AAC, nhưng bản thân `.mp4` chỉ là cái hộp chứa.",
    )

    h2(doc, "Ví dụ đọc pipeline từ một lệnh")
    code(doc, "ffmpeg -i input.mov -vf scale=1280:-2 -c:v libx264 -crf 23 -c:a aac output.mp4")
    bullets(
        doc,
        [
            "`-i input.mov`: đọc input.",
            "`-vf scale=1280:-2`: áp video filter, resize chiều rộng 1280 và tự tính chiều cao giữ tỉ lệ.",
            "`-c:v libx264`: encode video bằng H.264 encoder libx264.",
            "`-crf 23`: chọn mức chất lượng cho libx264.",
            "`-c:a aac`: encode audio thành AAC.",
            "`output.mp4`: muxer MP4 được chọn dựa theo đuôi file.",
        ],
    )

    h1(doc, "3. Cú pháp lệnh FFmpeg")
    p(doc, "Cú pháp tổng quát:")
    code(doc, "ffmpeg [global_options] [input_options] -i input [output_options] output")
    p(
        doc,
        "Điểm rất quan trọng: vị trí option quyết định nó áp dụng cho input hay output. Option đặt trước `-i` thường áp dụng cho input kế tiếp. Option đặt sau input và trước output thường áp dụng cho output đó.",
    )
    add_command_box(
        doc,
        "Convert cơ bản",
        "ffmpeg -i input.mp4 output.webm",
        [
            "FFmpeg tự đoán demuxer từ input và muxer từ đuôi output.",
            "FFmpeg tự chọn encoder phù hợp với output nếu bạn không chỉ định.",
            "Dễ dùng nhưng không nên phụ thuộc hoàn toàn trong production vì lựa chọn tự động có thể không đúng ý.",
        ],
    )
    add_command_box(
        doc,
        "Ghi đè output không hỏi",
        "ffmpeg -y -i input.mp4 output.mp4",
        [
            "`-y` là global option, cho phép ghi đè file output.",
            "Ngược lại, `-n` nghĩa là không ghi đè nếu file đã tồn tại.",
        ],
    )

    h1(doc, "4. Input, output và thứ tự option")
    p(
        doc,
        "FFmpeg có thể có nhiều input và nhiều output. Đây là điểm làm nó mạnh nhưng cũng dễ gây rối. Hãy đọc lệnh từ trái sang phải, mỗi khi thấy `-i` là một input mới, mỗi tên file/URL phía sau các output option là một output mới.",
    )
    code(doc, "ffmpeg -i video.mp4 -i logo.png -filter_complex \"[0:v][1:v]overlay=10:10\" -c:a copy out.mp4")
    bullets(
        doc,
        [
            "`video.mp4` là input 0.",
            "`logo.png` là input 1.",
            "`[0:v]` nghĩa là video stream từ input 0.",
            "`[1:v]` nghĩa là video stream từ input 1.",
            "`overlay=10:10` đặt logo lên video tại tọa độ x=10, y=10.",
        ],
    )

    h1(doc, "5. Stream là gì")
    p(
        doc,
        "Một file media thường không chỉ có một thứ. Nó có thể chứa nhiều stream: một video stream, nhiều audio stream, subtitle stream, metadata stream, attachment font... FFmpeg xử lý theo stream chứ không chỉ theo file.",
    )
    code(doc, "ffprobe -hide_banner -i input.mkv")
    p(doc, "Bạn có thể thấy output kiểu:")
    code(doc, "Stream #0:0: Video: h264...\nStream #0:1: Audio: aac, 48000 Hz, stereo...\nStream #0:2: Subtitle: subrip...")
    bullets(
        doc,
        [
            "`#0:0`: input số 0, stream số 0.",
            "`#0:1`: input số 0, stream số 1.",
            "`Video`, `Audio`, `Subtitle` là loại stream.",
            "Khi cần chọn stream chính xác, dùng `-map`.",
        ],
    )

    h1(doc, "6. Stream selection và -map")
    p(
        doc,
        "Nếu không dùng `-map`, FFmpeg tự chọn stream. Thường nó chọn video chất lượng cao nhất, audio phù hợp nhất và subtitle mặc định. Nhưng trong hệ thống thật, đặc biệt media monitoring hoặc file nhiều audio track, bạn nên dùng `-map` để tránh bất ngờ.",
    )
    add_command_box(
        doc,
        "Chọn video đầu tiên và audio thứ hai",
        "ffmpeg -i input.mkv -map 0:v:0 -map 0:a:1 -c copy output.mp4",
        [
            "`0:v:0`: input 0, video stream thứ 0.",
            "`0:a:1`: input 0, audio stream thứ 1.",
            "`-c copy`: không encode lại, chỉ copy stream đã chọn.",
        ],
    )
    add_command_box(
        doc,
        "Bỏ subtitle",
        "ffmpeg -i input.mkv -map 0:v -map 0:a -c copy output.mp4",
        [
            "Chỉ map video và audio.",
            "Subtitle không được đưa vào output.",
        ],
    )
    add_command_box(
        doc,
        "Map tất cả stream",
        "ffmpeg -i input.mkv -map 0 -c copy output.mkv",
        [
            "`-map 0` nghĩa là lấy tất cả stream từ input 0.",
            "Thường dùng khi remux sang MKV để giữ nguyên nhiều track.",
        ],
    )

    h1(doc, "7. Container, codec, encoder, decoder")
    p(doc, "Bốn khái niệm này phải tách bạch:")
    bullets(
        doc,
        [
            "Container: cái hộp chứa stream, ví dụ MP4, MKV, MOV, TS, WebM.",
            "Codec: chuẩn nén dữ liệu, ví dụ H.264, H.265/HEVC, AV1, AAC, Opus.",
            "Decoder: bộ giải mã codec để lấy dữ liệu thô.",
            "Encoder: bộ mã hóa dữ liệu thô thành codec.",
        ],
    )
    p(
        doc,
        "Một ví dụ đời thường: file `video.mp4` là container MP4, bên trong có video codec H.264 và audio codec AAC. Khi bạn đổi `.mkv` sang `.mp4` bằng `-c copy`, bạn chỉ đổi cái hộp nếu codec bên trong tương thích với MP4.",
    )
    add_command_box(
        doc,
        "Xem codec trong file",
        "ffprobe -v error -show_streams -show_format input.mp4",
        [
            "`show_streams` cho biết codec, resolution, duration, bitrate của từng stream.",
            "`show_format` cho biết container, duration tổng, size, bitrate tổng.",
        ],
    )

    h1(doc, "8. Streamcopy: khi nào dùng -c copy")
    p(
        doc,
        "Streamcopy nghĩa là FFmpeg không decode và không encode lại. Nó chỉ lấy stream đã có rồi đóng gói sang output. Đây là cách nhanh nhất và ít làm hỏng chất lượng nhất.",
    )
    add_command_box(
        doc,
        "Remux MKV sang MP4",
        "ffmpeg -i input.mkv -c copy output.mp4",
        [
            "Rất nhanh nếu codec trong MKV tương thích với MP4.",
            "Nếu input có subtitle hoặc audio codec MP4 không hỗ trợ, lệnh có thể lỗi.",
        ],
    )
    add_command_box(
        doc,
        "Cắt nhanh",
        "ffmpeg -ss 00:02:00 -i input.mp4 -t 30 -c copy clip.mp4",
        [
            "Nhanh vì không encode.",
            "Có thể không cắt chính xác tuyệt đối từng frame vì phụ thuộc keyframe.",
        ],
    )
    p(
        doc,
        "Dùng `-c copy` khi bạn không cần thay đổi nội dung hình/âm. Không dùng được nếu bạn cần resize, crop, chèn logo, đổi FPS, burn subtitle, normalize audio hoặc giảm bitrate bằng encode lại.",
    )

    h1(doc, "9. Transcoding: khi nào phải encode lại")
    p(
        doc,
        "Transcoding là decode rồi encode lại. Bạn cần transcode khi muốn thay đổi bản chất stream: đổi codec, giảm dung lượng, thay resolution, đổi FPS, áp filter hoặc làm output tương thích với trình phát cụ thể.",
    )
    add_command_box(
        doc,
        "Encode H.264 phổ thông",
        "ffmpeg -i input.mov -c:v libx264 -crf 23 -preset medium -c:a aac -b:a 128k output.mp4",
        [
            "`libx264`: encoder H.264 chất lượng tốt, rất phổ biến.",
            "`-crf 23`: chất lượng video, số thấp hơn nghĩa là đẹp hơn nhưng file lớn hơn.",
            "`-preset medium`: cân bằng tốc độ encode và hiệu quả nén.",
            "`-b:a 128k`: bitrate audio.",
        ],
    )
    h2(doc, "CRF nên hiểu thế nào")
    bullets(
        doc,
        [
            "CRF là chế độ giữ chất lượng tương đối ổn định, file size tự thay đổi theo độ phức tạp video.",
            "Với libx264, CRF 18 gần như rất đẹp, 23 là mặc định hợp lý, 28 nhỏ hơn nhưng xấu hơn.",
            "Với libx265, cùng cảm nhận chất lượng thường dùng CRF cao hơn libx264 một chút.",
            "Nếu bắt buộc bitrate cố định cho streaming, dùng bitrate/rate control thay vì chỉ CRF.",
        ],
    )

    h1(doc, "10. Bitrate, quality và preset")
    p(
        doc,
        "Ba thứ hay bị nhầm: bitrate là dung lượng dữ liệu mỗi giây, quality là cảm nhận hình ảnh/âm thanh, preset là mức độ encoder chịu tốn thời gian để nén hiệu quả hơn.",
    )
    bullets(
        doc,
        [
            "Cùng một bitrate, preset chậm hơn thường cho chất lượng tốt hơn.",
            "Cùng một CRF, preset chậm hơn thường cho file nhỏ hơn một chút.",
            "Preset không phải quality trực tiếp. `slow` không tự làm video đẹp hơn nếu bitrate/CRF không hợp lý.",
            "Với livestream, cần quan tâm bitrate tối đa, buffer và khả năng mạng.",
        ],
    )
    code(doc, "ffmpeg -i input.mp4 -c:v libx264 -b:v 3000k -maxrate 3500k -bufsize 7000k -c:a aac -b:a 128k out.mp4")

    h1(doc, "11. Filter: phần làm FFmpeg trở nên mạnh")
    p(
        doc,
        "Filter là nơi bạn biến đổi dữ liệu sau decode. Video filter xử lý frame hình ảnh. Audio filter xử lý sample âm thanh. Khi có filter, thường bạn không thể `-c copy` stream đó nữa vì dữ liệu đã bị thay đổi và phải encode lại.",
    )
    add_command_box(
        doc,
        "Resize video",
        "ffmpeg -i input.mp4 -vf scale=1280:-2 -c:v libx264 -c:a copy output.mp4",
        [
            "`scale=1280:-2`: đặt chiều rộng 1280, chiều cao tự tính và chia hết cho 2.",
            "Video phải encode lại vì đã resize.",
            "Audio có thể copy vì không đụng vào audio.",
        ],
    )
    add_command_box(
        doc,
        "Crop video",
        "ffmpeg -i input.mp4 -vf crop=1280:720:0:0 output.mp4",
        [
            "`crop=w:h:x:y`: cắt vùng có width, height, tọa độ bắt đầu x/y.",
            "Hữu ích khi bỏ viền đen hoặc lấy vùng màn hình.",
        ],
    )
    add_command_box(
        doc,
        "Normalize loudness audio",
        "ffmpeg -i input.mp4 -af loudnorm -c:v copy -c:a aac output.mp4",
        [
            "`loudnorm` xử lý âm lượng theo chuẩn loudness.",
            "Audio phải encode lại.",
            "Video copy được nếu không chỉnh hình.",
        ],
    )

    h1(doc, "12. Simple filtergraph và complex filtergraph")
    p(
        doc,
        "Simple filtergraph dùng khi xử lý một stream đơn giản, ví dụ `-vf scale=...` hoặc `-af volume=...`. Complex filtergraph dùng khi có nhiều input, nhiều output hoặc cần nối nhiều nhánh xử lý.",
    )
    code(doc, "ffmpeg -i input.mp4 -vf \"fps=30,scale=1280:-2\" output.mp4")
    p(doc, "Dấu phẩy nghĩa là filter chạy nối tiếp: đầu tiên đổi FPS, sau đó resize.")
    code(doc, "ffmpeg -i video.mp4 -i logo.png -filter_complex \"[0:v][1:v]overlay=W-w-20:20[v]\" -map \"[v]\" -map 0:a? output.mp4")
    bullets(
        doc,
        [
            "`[0:v]` và `[1:v]` là nhãn input stream.",
            "`overlay=W-w-20:20` đặt logo cách mép phải 20 pixel và mép trên 20 pixel.",
            "`[v]` là nhãn output của filtergraph.",
            "`0:a?` nghĩa là map audio nếu có, không lỗi nếu không có audio.",
        ],
    )

    h1(doc, "13. Time: -ss, -t, -to và vị trí đặt")
    p(
        doc,
        "Cắt video là tác vụ rất hay gặp, nhưng cần hiểu `-ss` trước hay sau `-i` có khác biệt. Đặt `-ss` trước input thường nhanh hơn vì seek ở input. Đặt sau input thường chính xác hơn vì decode tới vị trí đó.",
    )
    add_command_box(
        doc,
        "Cắt nhanh",
        "ffmpeg -ss 00:10:00 -i input.mp4 -t 60 -c copy output.mp4",
        [
            "Nhanh, phù hợp lấy đoạn gần đúng.",
            "Có thể lệch vài frame do keyframe.",
        ],
    )
    add_command_box(
        doc,
        "Cắt chính xác",
        "ffmpeg -i input.mp4 -ss 00:10:00 -t 60 -c:v libx264 -c:a aac output.mp4",
        [
            "Chính xác hơn vì decode rồi cắt.",
            "Chậm hơn vì encode lại.",
        ],
    )

    h1(doc, "14. HLS và media playlist/master playlist")
    p(
        doc,
        "HLS dùng file `.m3u8` làm playlist. Có hai khái niệm bạn đang gặp trong code Python: master playlist và media playlist.",
    )
    bullets(
        doc,
        [
            "Master playlist liệt kê nhiều variant, ví dụ 360p, 720p, 1080p, mỗi variant có bandwidth/resolution/codecs.",
            "Media playlist liệt kê các segment thực tế, thường là `.ts` hoặc `.m4s`.",
            "Trong thư viện Python `m3u8`, master thường có `playlist.is_variant == True` và có `playlist.playlists`.",
            "Media playlist thường có `playlist.segments` và không nên xử lý như master.",
        ],
    )
    h2(doc, "Tải HLS bằng FFmpeg")
    code(doc, "ffmpeg -i https://example.com/master.m3u8 -c copy output.mp4")
    bullets(
        doc,
        [
            "Nếu stream không mã hóa DRM và server cho phép, FFmpeg có thể đọc playlist và tải segment.",
            "`-c copy` giữ nguyên audio/video, chỉ đóng gói lại sang MP4.",
            "Nếu MP4 không tương thích với stream gốc, có thể cần transcode.",
        ],
    )
    h2(doc, "Tạo HLS từ file")
    code(doc, "ffmpeg -i input.mp4 -c:v libx264 -c:a aac -f hls -hls_time 6 -hls_list_size 0 playlist.m3u8")
    bullets(
        doc,
        [
            "`-f hls`: chọn HLS muxer.",
            "`-hls_time 6`: mỗi segment khoảng 6 giây.",
            "`-hls_list_size 0`: giữ toàn bộ segment trong playlist, phù hợp VOD.",
        ],
    )
    h2(doc, "Tạo nhiều chất lượng HLS")
    code(
        doc,
        "ffmpeg -i input.mp4 "
        "-filter_complex \"[0:v]split=2[v1][v2];[v1]scale=1280:-2[v1out];[v2]scale=854:-2[v2out]\" "
        "-map \"[v1out]\" -map 0:a -c:v:0 libx264 -b:v:0 3000k "
        "-map \"[v2out]\" -map 0:a -c:v:1 libx264 -b:v:1 1200k "
        "-c:a aac -f hls -var_stream_map \"v:0,a:0 v:1,a:1\" master.m3u8",
    )
    p(
        doc,
        "Lệnh nhiều variant thường dài và dễ sai. Trong production nên viết script hoặc dùng template, đồng thời kiểm tra output bằng ffprobe hoặc player HLS.",
    )

    h1(doc, "15. Livestream: RTMP, RTSP, SRT, UDP")
    p(
        doc,
        "FFmpeg không chỉ xử lý file tĩnh. Nó có thể đọc/ghi stream mạng. Khi làm live, ngoài codec còn phải quan tâm latency, buffer, reconnect, timestamp và độ ổn định mạng.",
    )
    add_command_box(
        doc,
        "Đẩy RTMP",
        "ffmpeg -re -i input.mp4 -c:v libx264 -c:a aac -f flv rtmp://server/app/stream_key",
        [
            "`-re` đọc input theo tốc độ thời gian thực, hữu ích khi giả lập live từ file.",
            "`-f flv` thường dùng cho RTMP output.",
        ],
    )
    add_command_box(
        doc,
        "Đọc RTSP camera",
        "ffmpeg -rtsp_transport tcp -i rtsp://user:pass@camera/stream -c copy output.mp4",
        [
            "`-rtsp_transport tcp` thường ổn định hơn UDP trong mạng không tốt.",
            "Nếu camera stream H.264/AAC hợp lệ, có thể copy để lưu nhanh.",
        ],
    )
    add_command_box(
        doc,
        "Gửi SRT",
        "ffmpeg -re -i input.mp4 -f mpegts \"srt://host:port?mode=caller&latency=200000\"",
        [
            "SRT phù hợp truyền live qua mạng có packet loss.",
            "Latency, mode caller/listener và passphrase là các option hay gặp.",
        ],
    )

    h1(doc, "16. ffprobe: công cụ bạn nên dùng trước khi viết lệnh")
    p(
        doc,
        "Nhiều lỗi FFmpeg đến từ việc đoán sai input. Trước khi xử lý, hãy dùng ffprobe để biết file có stream gì, codec nào, duration bao nhiêu, timebase thế nào.",
    )
    add_command_box(
        doc,
        "Xem tổng quan",
        "ffprobe -hide_banner -i input.mp4",
        [
            "Dễ đọc bằng mắt.",
            "Phù hợp debug nhanh.",
        ],
    )
    add_command_box(
        doc,
        "Xuất JSON để code xử lý",
        "ffprobe -v error -print_format json -show_format -show_streams input.mp4",
        [
            "Dùng tốt trong Python/Node/backend.",
            "Có thể parse codec_name, width, height, duration, bit_rate, tags.",
        ],
    )
    add_command_box(
        doc,
        "Chỉ lấy duration",
        "ffprobe -v error -show_entries format=duration -of default=nw=1:nk=1 input.mp4",
        [
            "Trả về số giây.",
            "Hữu ích trong pipeline batch processing.",
        ],
    )

    h1(doc, "17. Các nhóm option hay gặp")
    h2(doc, "Video")
    bullets(
        doc,
        [
            "`-c:v`: chọn video encoder hoặc `copy`.",
            "`-b:v`: đặt video bitrate.",
            "`-crf`: đặt chất lượng theo CRF cho một số encoder như libx264/libx265.",
            "`-preset`: tốc độ/hiệu quả encode.",
            "`-r`: đặt frame rate output hoặc input tùy vị trí.",
            "`-pix_fmt`: pixel format, ví dụ `yuv420p` để tương thích rộng.",
        ],
    )
    h2(doc, "Audio")
    bullets(
        doc,
        [
            "`-c:a`: chọn audio encoder hoặc `copy`.",
            "`-b:a`: audio bitrate.",
            "`-ar`: sample rate, ví dụ 48000.",
            "`-ac`: số kênh, ví dụ 2 cho stereo.",
            "`-af`: audio filter.",
        ],
    )
    h2(doc, "Subtitle")
    bullets(
        doc,
        [
            "`-c:s copy`: copy subtitle stream.",
            "`-sn`: bỏ subtitle.",
            "`subtitles=...`: burn subtitle vào video bằng filter.",
            "Subtitle mềm là stream riêng, subtitle burn-in là chữ đã dính vào hình.",
        ],
    )
    h2(doc, "Input/output control")
    bullets(
        doc,
        [
            "`-ss`: seek tới thời điểm.",
            "`-t`: duration cần lấy.",
            "`-to`: thời điểm kết thúc.",
            "`-y`: ghi đè output.",
            "`-hide_banner`: ẩn banner cho log gọn.",
            "`-loglevel`: điều chỉnh mức log.",
        ],
    )

    h1(doc, "18. Hardware acceleration")
    p(
        doc,
        "FFmpeg có thể dùng phần cứng để tăng tốc decode/encode, ví dụ NVIDIA NVENC/CUDA, Intel QSV, VAAPI trên Linux, VideoToolbox trên macOS. Nhưng phần cứng không tự nhiên làm mọi thứ tốt hơn. Nó giúp nhanh hơn, đôi khi chất lượng trên cùng bitrate kém hơn encoder phần mềm tốt như libx264 ở preset chậm.",
    )
    code(doc, "ffmpeg -hwaccels")
    bullets(
        doc,
        [
            "Dùng hardware khi cần tốc độ, realtime hoặc xử lý nhiều luồng.",
            "Dùng software encoder khi ưu tiên chất lượng/nén tốt và có thời gian encode.",
            "Cần kiểm tra build FFmpeg của bạn có encoder phần cứng hay không bằng `ffmpeg -encoders`.",
        ],
    )
    code(doc, "ffmpeg -i input.mp4 -c:v h264_nvenc -b:v 4000k -c:a aac output.mp4")

    h1(doc, "19. Debug lỗi FFmpeg theo kiểu có phương pháp")
    p(
        doc,
        "Khi FFmpeg lỗi, đừng đọc toàn bộ log từ đầu đến cuối ngay. Hãy tìm dòng lỗi gần cuối, rồi quay lên xem input/output context. Sau đó kiểm tra từng nhóm nguyên nhân.",
    )
    nums(
        doc,
        [
            "Chạy `ffprobe` để xác nhận input có đọc được không.",
            "Thêm `-hide_banner` để log gọn hơn.",
            "Thử `-t 5` để xử lý 5 giây đầu, tránh chờ lâu.",
            "Nếu lỗi codec/container, thử output MKV vì MKV dễ chứa nhiều codec hơn MP4.",
            "Nếu lỗi filtergraph, tách lệnh thành từng filter nhỏ.",
            "Nếu lỗi network, kiểm tra timeout, reconnect, headers, protocol option và quyền truy cập URL.",
        ],
    )
    h2(doc, "Một số lỗi thường gặp")
    bullets(
        doc,
        [
            "`Unknown encoder`: bản FFmpeg không có encoder đó hoặc viết sai tên.",
            "`Invalid data found when processing input`: input không đọc được, URL sai, file hỏng hoặc format không đúng.",
            "`Could not write header`: output container không hỗ trợ stream/codec bạn đang đưa vào.",
            "`Non-monotonous DTS`: vấn đề timestamp, thường gặp khi remux hoặc nối file.",
            "`Protocol not found`: FFmpeg build thiếu protocol hoặc URL scheme sai.",
        ],
    )

    h1(doc, "20. Công thức thực hành theo nhu cầu")
    add_command_box(
        doc,
        "Tạo thumbnail tại giây thứ 10",
        "ffmpeg -ss 10 -i input.mp4 -frames:v 1 thumbnail.jpg",
        ["Dùng cho preview video hoặc media monitoring dashboard."],
    )
    add_command_box(
        doc,
        "Tách audio WAV",
        "ffmpeg -i input.mp4 -vn -acodec pcm_s16le -ar 16000 -ac 1 audio.wav",
        ["Hay dùng trước khi đưa audio vào speech-to-text hoặc phân tích âm thanh."],
    )
    add_command_box(
        doc,
        "Ghép nhiều file cùng codec",
        "ffmpeg -f concat -safe 0 -i list.txt -c copy joined.mp4",
        ["File `list.txt` chứa từng dòng dạng: `file 'part1.mp4'`."],
    )
    add_command_box(
        doc,
        "Convert sang MP4 tương thích rộng",
        "ffmpeg -i input.mkv -c:v libx264 -pix_fmt yuv420p -c:a aac -movflags +faststart output.mp4",
        [
            "`yuv420p` giúp tương thích nhiều trình phát.",
            "`+faststart` đưa metadata MP4 lên đầu file, tốt cho phát trên web.",
        ],
    )
    add_command_box(
        doc,
        "Giảm dung lượng video",
        "ffmpeg -i input.mp4 -c:v libx264 -crf 28 -preset medium -c:a aac -b:a 96k small.mp4",
        [
            "Tăng CRF để giảm dung lượng.",
            "Đừng tăng quá cao nếu cần giữ chi tiết hình ảnh.",
        ],
    )
    add_command_box(
        doc,
        "Ghi màn hình Windows bằng gdigrab",
        "ffmpeg -f gdigrab -framerate 30 -i desktop output.mp4",
        ["Phù hợp test nhanh, nhưng tùy bản FFmpeg và Windows có thể cần chỉnh thêm audio input."],
    )

    h1(doc, "21. Cách học FFmpeg hiệu quả")
    p(
        doc,
        "Documentation của FFmpeg rất lớn vì nó gom gần như toàn bộ option của nhiều thư viện, codec, filter, format và protocol. Cách học tốt không phải đọc từ đầu đến cuối, mà là đi từ pipeline và use case.",
    )
    nums(
        doc,
        [
            "Học cú pháp chung và vị trí option.",
            "Hiểu stream/container/codec/encoder/decoder.",
            "Dùng ffprobe để nhìn file.",
            "Tập streamcopy và transcode cơ bản.",
            "Học filter đơn giản: scale, crop, fps, volume.",
            "Học `-map` khi file có nhiều stream.",
            "Sau đó mới đi vào HLS, RTSP, SRT, hardware acceleration hoặc filtergraph phức tạp.",
        ],
    )
    p(
        doc,
        "Khi gặp một option trong docs gốc, hãy tự hỏi: option này thuộc giai đoạn nào của pipeline? Nó áp dụng cho input, decoder, filter, encoder, muxer, protocol hay output? Câu hỏi đó giúp bạn không bị lạc trong hàng nghìn mục documentation.",
    )

    h1(doc, "22. Liên hệ với code Python m3u8 của bạn")
    p(
        doc,
        "Trong project hiện tại bạn đang dùng thư viện Python `m3u8` để inspect playlist. FFmpeg cũng đọc được HLS trực tiếp, nhưng Python `m3u8` giúp bạn hiểu cấu trúc playlist trước khi quyết định làm gì tiếp.",
    )
    bullets(
        doc,
        [
            "Nếu URL là master playlist, bạn đọc `playlist.playlists` để lấy các variant.",
            "Mỗi variant có `stream_info`: bandwidth, resolution, codecs.",
            "Nếu URL là media playlist, bạn đọc `playlist.segments` để lấy danh sách segment.",
            "Đừng xử lý media playlist như master playlist vì media playlist không có danh sách variant.",
            "Sau khi chọn variant hoặc segment, bạn có thể dùng FFmpeg để tải, remux, transcode hoặc tạo preview.",
        ],
    )
    code(doc, "ffmpeg -i \"https://example.com/master.m3u8\" -c copy output.mp4")
    p(
        doc,
        "Trong media monitoring, một flow thực tế là: Python kiểm tra playlist và metadata, FFmpeg xử lý media, ffprobe xác nhận output, sau đó hệ thống lưu kết quả/log/trạng thái.",
    )

    h1(doc, "23. Mini glossary")
    bullets(
        doc,
        [
            "Demux: tách container thành các stream.",
            "Mux: đóng gói stream vào container.",
            "Remux: đổi container nhưng không đổi codec.",
            "Transcode: decode rồi encode lại.",
            "Bitrate: số bit mỗi giây.",
            "Frame rate/FPS: số frame video mỗi giây.",
            "Sample rate: số mẫu audio mỗi giây.",
            "Keyframe: frame độc lập, quan trọng khi seek/cắt nhanh.",
            "Filtergraph: sơ đồ các filter nối với nhau.",
            "HLS: HTTP Live Streaming, dùng playlist m3u8 và segment.",
        ],
    )

    h1(doc, "24. Nguồn chính thức nên đọc tiếp")
    bullets(
        doc,
        [
            f"FFmpeg all-in-one documentation: {SOURCE_URL}",
            "Documentation index: https://ffmpeg.org/documentation.html",
            "ffmpeg command manual: https://ffmpeg.org/ffmpeg.html",
            "ffprobe manual: https://ffmpeg.org/ffprobe.html",
            "Filters manual: https://ffmpeg.org/ffmpeg-filters.html",
            "Formats manual: https://ffmpeg.org/ffmpeg-formats.html",
            "Protocols manual: https://ffmpeg.org/ffmpeg-protocols.html",
        ],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
    print(OUT_PATH.resolve())
